#!/usr/bin/env python3
"""
JEEVAN-AIR | Prototype and Workflow DOCX Generator
Team ZYNTAX — SIH26177 (Qualcomm Inc)
Builds Deliverable 3: JEEVAN-AIR_Prototype_and_Workflow.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

doc_path = "/Users/sheshagiri/jeevan air/JEEVAN-AIR_Prototype_and_Workflow.docx"
screenshots_dir = "/Users/sheshagiri/jeevan air/screenshots"

doc = Document()

# Set standard page margins (1 inch)
sections = doc.sections
for s in sections:
    s.top_margin = Inches(0.9)
    s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(0.9)
    s.right_margin = Inches(0.9)

# Colors
COLOR_NAVY = RGBColor(15, 23, 42)
COLOR_CYAN = RGBColor(2, 132, 199)
COLOR_SLATE = RGBColor(100, 116, 139)
COLOR_GREEN = RGBColor(21, 128, 61)
COLOR_AMBER = RGBColor(180, 83, 9)

def set_cell_shading(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def add_styled_heading(text, level, color=COLOR_NAVY):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def add_callout(text, title="DISCLAIMER", border_color="f59e0b", bg_color="fffbeb", text_color="92400e"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.7)
    set_cell_shading(cell, bg_color)
    
    # Border
    tcPr = cell._tc.get_or_add_tcPr()
    borders_xml = f'''
    <w:tcBorders {nsdecls("w")}>
        <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>
        <w:top w:val="none"/>
        <w:right w:val="none"/>
        <w:bottom w:val="none"/>
    </w:tcBorders>
    '''
    tcPr.append(parse_xml(borders_xml))
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r_title = p.add_run(f"[{title}] ")
    r_title.bold = True
    r_title.font.size = Pt(9.5)
    r_title.font.color.rgb = COLOR_AMBER
    
    r_text = p.add_run(text)
    r_text.font.size = Pt(9)
    r_text.font.color.rgb = RGBColor(146, 64, 14)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_feature_box(number, title, purpose, what_user_sees, how_software_handles, status_tag):
    h = add_styled_heading(f"{number}. {title}", level=2, color=COLOR_CYAN)
    
    p_status = doc.add_paragraph()
    p_status.paragraph_format.space_after = Pt(4)
    r_tag = p_status.add_run(f"CLASSIFICATION: [{status_tag}]")
    r_tag.bold = True
    if "IMPLEMENTED" in status_tag:
        r_tag.font.color.rgb = COLOR_GREEN
    elif "SIMULATED" in status_tag:
        r_tag.font.color.rgb = COLOR_CYAN
    else:
        r_tag.font.color.rgb = COLOR_AMBER
        
    p_purp = doc.add_paragraph()
    p_purp.paragraph_format.space_after = Pt(3)
    p_purp.add_run("• Purpose: ").bold = True
    p_purp.add_run(purpose)
    
    p_sees = doc.add_paragraph()
    p_sees.paragraph_format.space_after = Pt(3)
    p_sees.add_run("• What the User Sees: ").bold = True
    p_sees.add_run(what_user_sees)
    
    p_hand = doc.add_paragraph()
    p_hand.paragraph_format.space_after = Pt(6)
    p_hand.add_run("• How the Software Handles It: ").bold = True
    p_hand.add_run(how_software_handles)

def add_image_box(img_path, caption, label="SOFTWARE PROTOTYPE"):
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(2)
        doc.add_picture(img_path, width=Inches(5.8))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        r_lbl = p_cap.add_run(f"[{label}] ")
        r_lbl.bold = True
        r_lbl.font.size = Pt(8.5)
        r_lbl.font.color.rgb = COLOR_CYAN
        
        r_cap = p_cap.add_run(caption)
        r_cap.font.size = Pt(8.5)
        r_cap.italic = True
        r_cap.font.color.rgb = COLOR_SLATE

print("Writing document title and metadata...")

# Title Header
p_top = doc.add_paragraph()
p_top.paragraph_format.space_before = Pt(10)
p_top.paragraph_format.space_after = Pt(2)
r_pre = p_top.add_run("SMART INDIA HACKATHON 2026 • PROBLEM STATEMENT SIH26177")
r_pre.font.size = Pt(9)
r_pre.font.bold = True
r_pre.font.color.rgb = COLOR_CYAN

p_title = doc.add_paragraph()
p_title.paragraph_format.space_after = Pt(2)
r_main = p_title.add_run("JEEVAN-AIR")
r_main.font.size = Pt(26)
r_main.font.bold = True
r_main.font.color.rgb = COLOR_NAVY

p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_after = Pt(8)
r_sub = p_sub.add_run("Aerial Intelligence & Rescue — Prototype Specification & System Workflow Guide")
r_sub.font.size = Pt(13)
r_sub.font.color.rgb = COLOR_SLATE

# Meta Table
tbl_meta = doc.add_table(rows=4, cols=2)
tbl_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_meta.autofit = False
meta_data = [
    ("Team", "TEAM ZYNTAX"),
    ("Problem Statement / Sponsor", "SIH26177 / Qualcomm Inc"),
    ("Category & Theme", "Hardware / Robotics and Drones"),
    ("Deliverable Type", "Software Prototype Technical Specification & System Workflow Guide")
]
for idx, (lbl, val) in enumerate(meta_data):
    cell_lbl = tbl_meta.cell(idx, 0)
    cell_val = tbl_meta.cell(idx, 1)
    cell_lbl.width = Inches(2.2)
    cell_val.width = Inches(4.5)
    set_cell_shading(cell_lbl, "f1f5f9")
    set_cell_shading(cell_val, "ffffff")
    
    p0 = cell_lbl.paragraphs[0]
    p0.paragraph_format.space_before = Pt(3)
    p0.paragraph_format.space_after = Pt(3)
    r0 = p0.add_run(lbl)
    r0.bold = True
    r0.font.size = Pt(9)
    
    p1 = cell_val.paragraphs[0]
    p1.paragraph_format.space_before = Pt(3)
    p1.paragraph_format.space_after = Pt(3)
    r1 = p1.add_run(val)
    r1.font.size = Pt(9)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

add_callout(
    "PHYSICAL HARDWARE DISCLAIMER: The physical drone does NOT exist yet. The current JEEVAN-AIR deliverable is a functional Ground Control Station (GCS) software and rescue intelligence prototype. All drone telemetry, GPS positions, optical camera feeds, and thermal readings are mathematically simulated in software. No simulated data is described as real hardware data.",
    title="IMPORTANT CURRENT STATUS"
)

# ==============================================================================
# PART A — SOFTWARE PROTOTYPE
# ==============================================================================
add_styled_heading("PART A — SOFTWARE PROTOTYPE", level=1, color=COLOR_NAVY)

add_feature_box(
    1, "Prototype Overview",
    "Provide search-and-rescue mission commanders with an autonomous aerial command interface that transforms raw sensor streams into prioritized, explainable victim triage and safe ground access corridors.",
    "A responsive, dark-themed tactical dashboard featuring 11 navigation modules (Dashboard, Mission Control, Live Search, Detections, Hazards, Search Map, Alerts, Mission History, System Status, System Architecture, and HW Integration).",
    "Executes as a single-page React 18 / TypeScript application on Vite (Port 5180), maintaining reactive state synced to the SimulationDataProvider via an event bus.",
    "IMPLEMENTED PROTOTYPE"
)

add_feature_box(
    2, "Current Prototype Status",
    "Establish verified software and decision-making logic before investing in physical drone assembly and flight controller wiring.",
    "Live dashboard with active telemetry counters, simulated 3D GPS lock badges, synthetic camera overlays, real-time alert notifications, and multi-modal AI benchmark counters.",
    "All flight movements step along a 10-waypoint lawn-mower route across sectors A1–C3. Incident injections simulate victim detections, second-look reviews, and hazard discoveries in real time.",
    "IMPLEMENTED (SOFTWARE) / SIMULATED (DATA)"
)

add_feature_box(
    3, "Dashboard Overview",
    "Surface critical operational intelligence on a single high-contrast screen without cognitive overload during high-stress rescue sorties.",
    "Top mission statistics (surveyed sectors, identified survivors, elapsed time), active mission objective banner, controls bar, camera feed viewport, tactical sector map, drone telemetry card, alerts stream, AI diagnostic panel, and rescue priority queue.",
    "MissionContext.tsx subscribes to the active data provider (createDataProvider() factory), updating all child components whenever position steps, new victims are localized, or hazard buffers expand.",
    "IMPLEMENTED"
)

add_image_box(f"{screenshots_dir}/sc_header_stats.png", "Header & Mission Progress Stat Cards")

add_feature_box(
    4, "Main Dashboard View",
    "Consolidate drone status, camera imagery, and ground safe corridors in a synchronized dual-column operational layout.",
    "Left 7-column layout displaying the live camera feed with HUD overlays and the 3x3 tactical GIS map. Right 5-column layout displaying drone flight gauges and the real-time alert feed.",
    "Rendered via DashboardPage.tsx. Adapts smoothly to desktop and field command tablet displays.",
    "IMPLEMENTED"
)

add_image_box(f"{screenshots_dir}/dashboard.png", "Main Ground Control Station Command Center (Full View)")

add_feature_box(
    5, "Simulated Drone Telemetry",
    "Model the physical dynamics and flight performance of a medium-lift SAR quadcopter.",
    "Drone ID JA-RESCUE-01, Altitude (32m AGL), Ground Speed (0 to 3.5 m/s), Heading (45 deg NE), Battery (98% discharging at 0.05%/s), Estimated Flight Time (~22 min), and Comms Link (94% RSSI).",
    "Calculated in SimulationDataProvider.ts using kinematic position interpolation between waypoints. Battery decrements based on motor hover and forward transit energy profiles.",
    "SIMULATED DATA"
)

add_image_box(f"{screenshots_dir}/sc_drone_telemetry.png", "Simulated Drone Telemetry Gauges & Communication State")

add_feature_box(
    6, "Simulated GPS & Satellite Fix",
    "Simulate high-precision geospatial localization in a realistic disaster operations theater.",
    "GPS Fix badge: SIMULATED — LOCKED (3D FIX, 14 Satellites, HDOP 1.1). Coordinates centered on Chennai disaster zone (13.0827 deg N, 80.2707 deg E) across sectors A1 to C3.",
    "Generates WGS-84 decimal degree coordinates matching the drone's position on the sector grid. Formatted to conform to future u-blox NEO-M9N GNSS NMEA/UBX payloads.",
    "SIMULATED DATA"
)

add_feature_box(
    7, "Survivor Detection Interface",
    "Localize individuals on the ground using optical silhouette detection reinforced by thermal body heat verification.",
    "Multi-spectral bounding boxes over victim positions, optical confidence (e.g. 94%), thermal confidence (e.g. 91%), biometric temperature readout (e.g. 36.8 C), condition badges (CRITICAL, STABLE, UNCERTAIN), and priority ranks (#1, #2).",
    "YoloVisionProvider extracts candidate silhouettes; ThermalProcessingProvider evaluates radiometric heat matrices (35.5 C to 38.2 C); MultiModalFusionEngine correlates both using spatial IoU (>= 0.20 threshold).",
    "IMPLEMENTED (PIPELINE) / SIMULATED (FRAMES)"
)

add_image_box(f"{screenshots_dir}/sc_camera_feed.png", "Multi-Spectral Camera Feed Viewport (Optical, Thermal & AI Overlays)")

add_feature_box(
    8, "Hazard Detection Interface",
    "Identify environmental dangers that threaten victims or block ground rescue personnel.",
    "Pulsing red circular danger rings on the tactical map showing hazard positions and radii (35m to 50m). Detailed hazard ledger on HazardsPage.tsx tracking Fire, Flood, Smoke, Debris, and Chemical Leaks.",
    "HazardDetectionProvider.ts classifies threats based on thermal combustion thresholds (> 60 C for fire), optical water reflection heuristics (flood), and surface roughness variances (debris).",
    "IMPLEMENTED (INTERFACE) / SIMULATED (INCIDENTS)"
)

add_feature_box(
    9, "Explainable Risk Assessment",
    "Calculate an objective, non-medical urgency score (0–100) with explicit human-readable reasons.",
    "Explainable Decision Card displaying Risk Score (e.g. 88 / 100), Severity Level (CRITICAL), and bulleted reasons: 'No movement detected: high risk of entrapment', 'Extreme proximity to critical threat: Fire (~32m)'.",
    "rescueIntelligence.ts evaluates 5 weighted factors: Movement (0–30 pts), Condition (0–25 pts), Thermal Heat Confirmation (0–15 pts), Hazard Proximity (0–20 pts), and Terrain Accessibility (0–10 pts).",
    "IMPLEMENTED"
)

add_feature_box(
    10, "Rescue Prioritization Queue",
    "Deliver an actionable triage queue so rescue commanders know exactly which survivor to extract first.",
    "Rescue Priority Panel showing survivors ranked sequentially: Priority #1 (Most Urgent), Priority #2, Priority #3... with priority badges on map icons and detection modals.",
    "prioritizeSurvivors() filters out rejected false positives, sorts active candidates by descending riskScore, resolves ties via sensor confidence, and assigns ordinal ranks.",
    "IMPLEMENTED"
)

add_feature_box(
    11, "Autonomous Second-Look Verification",
    "Eliminate optical false positives (shadows, blankets, debris) before ground resources are dispatched.",
    "Ambiguous detections marked 'VERIFICATION NEEDED (POSSIBLE)' with an active 'REQUEST SECOND LOOK' button. Status updates to 'IN PROGRESS' and resolves to either 'VERIFIED' (green) or 'REJECTED' (grayed out).",
    "executeSecondLookResolution() simulates multi-spectral re-examination: confirms candidates if optical confidence >= 65% or thermal heat is detected (boosting confidence >= 88%), or rejects false positives.",
    "IMPLEMENTED"
)

add_feature_box(
    12, "Safe-Access Ground Guidance",
    "Synthesize safe foot navigation paths for ground rescue teams avoiding deadly hazard sectors.",
    "Tactical Map displays an emerald dashed corridor with step markers ('Ingress Step #1', 'Step #2'...). Safe Route Card displays distance (e.g. 280m), walking time (4.2 min @ 1.1 m/s), avoided hazards, and terrain warnings.",
    "calculateSafeRoute() runs A* pathfinding over the 3x3 sector graph, marking sectors with active critical hazards as impassable obstacles and calculating the shortest safe corridor from base (A1).",
    "IMPLEMENTED"
)

add_image_box(f"{screenshots_dir}/sc_tactical_map.png", "Tactical 3x3 GIS Sector Map with Drone Flight Track & Ground Safe Corridor")

add_feature_box(
    13, "Dynamic Mission Replanning",
    "Automatically adapt drone search objectives upon discovering high-risk victims or advancing threats.",
    "Mission Objective Banner updates from 'AUTONOMOUS SWEEP: Sector Grid A1–C3' to 'PRIORITY REPLAN: Monitor & guide responder approach to Survivor in Sector B2'. Timeline logs the replanning event.",
    "SimulationDataProvider detects critical incidents and overrides the default sweep goal, logging chronological audit entries to DecisionTimeline.tsx.",
    "IMPLEMENTED"
)

add_feature_box(
    14, "Real-Time Alerts Feed",
    "Alert commanders immediately to high-priority events requiring tactical action.",
    "Alert cards tagged with severity (CRITICAL, HIGH, MEDIUM, LOW), title, message, and quick-action buttons ('VERIFY', 'DISMISS'). Unverified badge counter in the sidebar.",
    "Alerts are published whenever a survivor is localized, risk is elevated, or a hazard expands. Tracks lifecycle from UNREAD to REQUIRES_VERIFICATION to VERIFIED.",
    "IMPLEMENTED"
)

add_image_box(f"{screenshots_dir}/sc_alerts_panel.png", "Real-Time Incident Alert Feed with Urgency Badges")

add_feature_box(
    15, "Disaster Geospatial Map",
    "Provide complete situational awareness across the 3x3 operational grid.",
    "Interactive GIS grid showing drone flight track (cyan), drone position icon with heading, detected survivors (orange icons with rank badges), active hazards (red danger rings), and ground safe corridor (emerald).",
    "Rendered via SearchMap.tsx with SVG path overlays, sector centroid markers, and click-to-inspect sector inspectors.",
    "IMPLEMENTED"
)

add_feature_box(
    16, "Drone Status Display",
    "Display flight health and controller readiness for the operator.",
    "Flight Mode (AUTONOMOUS), Flight Controller (NOT CONNECTED — SIMULATED ROUTE), AI Status (ACTIVE), Battery gauge, Altitude AGL, Speed, Heading dial, and Comms Status.",
    "DroneStatusCard.tsx renders live gauges based on the DroneStatus model. Clearly flags that flight control is software-simulated.",
    "IMPLEMENTED"
)

add_feature_box(
    17, "Connectivity Status",
    "Simulate radio datalink health and emergency fail-safe behaviors.",
    "Header and telemetry card show 'SIMULATED — CONNECTED', 'DEGRADED', or 'DISCONNECTED'. Degraded mode warns of simulated packet latency; disconnected mode triggers RTL advisory.",
    "setConnectivity() method on SimulationDataProvider emulates signal degradation and datalink drops for command resilience testing.",
    "SIMULATED"
)

add_feature_box(
    18, "Current Technology Stack",
    "Build on modern, maintainable, strictly typed web technologies.",
    "React 18.3.1, TypeScript 5.6.3, Tailwind CSS 3.4.14, Vite 5.4.8, Lucide React 1.16.0, Node.js v22.6.0 on macOS Darwin ARM64. Port 5180.",
    "Strict type safety with isolatedModules enabled. Production build compiles clean in 1.4s with 355 kB bundle size.",
    "IMPLEMENTED"
)

add_feature_box(
    19, "Current Limitations",
    "Explicitly acknowledge prototype boundaries in accordance with hackathon evaluation standards.",
    "1. Physical drone airframe and motors absent; 2. Optical/thermal cameras absent; 3. Live MAVLink radio link absent; 4. Coarse 3x3 sector grid; 5. In-browser CPU inference.",
    "All software is engineered with clean abstraction boundaries so that hardware procurement in Phase 4 requires zero UI rewrites.",
    "CURRENT LIMITATION"
)

# ==============================================================================
# PART B — SYSTEM WORKFLOW
# ==============================================================================
doc.add_page_break()
add_styled_heading("PART B — SYSTEM WORKFLOW", level=1, color=COLOR_NAVY)

p_wf_intro = doc.add_paragraph()
p_wf_intro.add_run("The complete operational intelligence workflow of JEEVAN-AIR bridges raw aerial sensing and ground rescue execution:").bold = True

p_flow_str = doc.add_paragraph()
p_flow_str.paragraph_format.space_before = Pt(6)
p_flow_str.paragraph_format.space_after = Pt(10)
flow_steps = [
    "DISASTER EVENT", "DRONE DEPLOYMENT", "AREA SEARCH", "RGB + THERMAL INPUT",
    "AI DETECTION", "VERIFY (2ND-LOOK)", "HAZARD CONTEXT", "RISK ASSESSMENT",
    "RESCUE PRIORITY", "SAFE-ACCESS GUIDANCE", "RESPONDER ALERT", "MISSION UPDATE"
]
p_flow_str.add_run(" &longrightarrow; ".join(flow_steps))
p_flow_str.runs[0].font.size = Pt(8.5)
p_flow_str.runs[0].font.bold = True
p_flow_str.runs[0].font.color.rgb = COLOR_CYAN

add_styled_heading("Current Software Workflow vs. Future Hardware Workflow", level=2, color=COLOR_NAVY)

tbl_wf_comp = doc.add_table(rows=2, cols=2)
tbl_wf_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_wf_comp.autofit = False

c00 = tbl_wf_comp.cell(0, 0)
c01 = tbl_wf_comp.cell(0, 1)
c10 = tbl_wf_comp.cell(1, 0)
c11 = tbl_wf_comp.cell(1, 1)

c00.width = Inches(3.3)
c01.width = Inches(3.4)
c10.width = Inches(3.3)
c11.width = Inches(3.4)

set_cell_shading(c00, "0f172a")
set_cell_shading(c01, "0f172a")
set_cell_shading(c10, "f8fafc")
set_cell_shading(c11, "f8fafc")

p_h0 = c00.paragraphs[0]
r_h0 = p_h0.add_run("CURRENT SOFTWARE WORKFLOW [PROTOTYPE]")
r_h0.bold = True
r_h0.font.color.rgb = RGBColor(56, 189, 248)
r_h0.font.size = Pt(8.5)

p_h1 = c01.paragraphs[0]
r_h1 = p_h1.add_run("FUTURE PHYSICAL DRONE WORKFLOW [PLANNED]")
r_h1.bold = True
r_h1.font.color.rgb = RGBColor(251, 191, 36)
r_h1.font.size = Pt(8.5)

p_b0 = c10.paragraphs[0]
p_b0.add_run(
    "1. SIMULATED DATA GENERATOR\n"
    "   (SimulationDataProvider.ts)\n"
    "   ↓\n"
    "2. IDataAdapter INTERFACE CONTRACT\n"
    "   (providerFactory.ts)\n"
    "   ↓\n"
    "3. COMMON DATA MODEL\n"
    "   (types/common.ts)\n"
    "   ↓\n"
    "4. RESCUE INTELLIGENCE ENGINE\n"
    "   (Risk, Priority, Safe Route)\n"
    "   ↓\n"
    "5. MISSION CONTEXT & GCS DASHBOARD\n"
    "   (React 18 / Tailwind CSS)"
)
p_b0.runs[0].font.size = Pt(8)
p_b0.runs[0].font.name = "Courier"

p_b1 = c11.paragraphs[0]
p_b1.add_run(
    "1. PHYSICAL DRONE AIRFRAME\n"
    "   (Pixhawk 6C Autopilot / MAVLink 2.0)\n"
    "   ↓\n"
    "2. SENSORS (Sony RGB + FLIR LWIR)\n"
    "   ↓\n"
    "3. ONBOARD EDGE COMPUTE\n"
    "   (NVIDIA Jetson Orin Nano / YOLOv8)\n"
    "   ↓\n"
    "4. JEEVAN-AIR EDGE BRIDGE\n"
    "   (Python FastAPI WebSocket @ 10 Hz)\n"
    "   ↓\n"
    "5. HardwareDataProvider.ts\n"
    "   ↓\n"
    "6. GCS DASHBOARD (Unchanged UI)"
)
p_b1.runs[0].font.size = Pt(8)
p_b1.runs[0].font.name = "Courier"

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# 8 Workflow Descriptions
workflows = [
    ("1. End-to-End System Workflow",
     "Disaster Incident Detected -> GCS Commands Autonomous Search -> UAV Sweeps Sectors A1-C3 -> Dual-Camera Video Streams -> Edge AI Infers Humans & Hazards -> Multi-Modal Fusion Engine -> Rescue Intelligence Evaluates Risk -> Priority Queue Ranked -> Safe Ingress Corridor Synthesized -> Ground Team Alerted."),
    ("2. Survivor Detection Workflow",
     "RGB Optical Silhouette Extracted -> Thermal LWIR Radiometric Heat Measured -> Spatial IoU Bounding Box Overlap Calculated -> Case A (Dual Match: Fused Confidence Boosted, Status VERIFIED) | Case B (Optical Only: Discounted, Condition UNCERTAIN, Triggers 2nd-Look) | Case C (Thermal Only: Rubble Occlusion, Condition CRITICAL)."),
    ("3. Hazard Detection Workflow",
     "Thermal Hotspot Scan (> 60 C) / Optical Edge Contrast Analysis -> Threat Classified (Fire, Flood, Smoke, Debris) -> GPS Centroid & Danger Radius (35m-50m) Calculated -> Spatial Hazard Buffer Established -> Threat Added to Disaster GIS Map & Alert Stream."),
    ("4. Rescue Priority Workflow",
     "All Active Survivors Gathered -> Filter Out False-Positive Rejected Targets -> Compute Multi-Factor Risk Score (0-100) -> Primary Sort Descending Risk Score -> Secondary Sort Descending Fused Confidence -> Assign Ordinal Priority Ranks (#1, #2, #3) -> Update Dashboard Priority Queue."),
    ("5. Second-Look Verification Workflow",
     "Target Detected with Low Confidence (< 75%) or Unconfirmed Thermal Heat -> Verification Status set to POSSIBLE -> Operator or Autonomous Logic Clicks REQUEST SECOND LOOK -> Status set to IN_PROGRESS -> Drone Retargets / Executes Multi-Spectral Re-examination -> Status upgraded to VERIFIED or relegated to REJECTED."),
    ("6. Safe-Access Guidance Workflow",
     "Victim Localized in Target Sector -> Load 3x3 Sector Graph -> Mark Sectors with Active Critical Threats (Fire/Chemical) as IMPASSABLE Obstacles -> Apply Travel Cost Penalties to High Hazard Sectors -> Run A* Shortest Path from Staging Base (A1) -> Generate Step-by-Step Waypoints, Distance & Walking Time @ 1.1 m/s."),
    ("7. Dynamic Mission Replanning Workflow",
     "Drone Executing Standard Lawn-Mower Sweep -> Incident Event Triggered (CRITICAL Survivor / Fire Proximity) -> Mission Replanning Engine Interrupts Default Objective -> Dynamic Mission Banner Updates -> Flight Route Retargeted to Target Sector -> Chronological Event Logged to Decision Timeline."),
    ("8. Future Hardware Integration Workflow",
     "Procure Pixhawk 6C, u-blox GNSS, Jetson Orin Nano & FLIR Boson -> Connect Pixhawk UART to Jetson -> Deploy JeevanAir Edge Bridge (FastAPI) -> In src/hardware/config.ts set SIMULATION_MODE = false -> providerFactory.ts Instantiates HardwareDataProvider -> Dashboard Operates on Real Drone with Zero UI Modifications.")
]

for title, desc in workflows:
    h = add_styled_heading(title, level=3, color=COLOR_CYAN)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(desc)

# Embed ER Diagram in Part B
add_styled_heading("Logical ER Data Model Architecture", level=2, color=COLOR_NAVY)
p_er_desc = doc.add_paragraph()
p_er_desc.add_run("The logical entity-relationship architecture below represents all software data structures currently active in the prototype and provides the formal schema for future persistent storage (PostgreSQL/PostGIS):")

add_image_box("/Users/sheshagiri/jeevan air/JEEVAN-AIR_ER_Diagram.png", "JEEVAN-AIR Logical ER Data Model — 15 Entities Across 5 Pillars", label="LOGICAL DATA MODEL")

# AI Diagnostics Panel Screenshot
add_image_box(f"{screenshots_dir}/sc_ai_panel.png", "Live AI Inference & Multi-Modal Fusion Diagnostic Panel with Measured Latency", label="SOFTWARE PROTOTYPE")

# Summary Table
add_styled_heading("Summary: System Workflow Validation Status", level=2, color=COLOR_NAVY)
tbl_val = doc.add_table(rows=9, cols=3)
tbl_val.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_val.autofit = False

headers = ["Workflow Module", "Validation Mechanism", "Current Engineering Status"]
for col_idx, htext in enumerate(headers):
    cell = tbl_val.cell(0, col_idx)
    set_cell_shading(cell, "0f172a")
    p = cell.paragraphs[0]
    r = p.add_run(htext)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(8.5)

val_rows = [
    ("End-to-End Decision Chain", "Automated regression suite (test:all)", "85 / 85 PASSED (100%)"),
    ("Multi-Factor Risk Scoring", "24 assertions in rescueIntelligence.test.ts", "VERIFIED IN CODE"),
    ("Priority Queue Ranking", "Descending risk sort with confidence tie-breaker", "VERIFIED IN CODE"),
    ("A* Safe Ground Routing", "Graph pathfinder bypassing active fire hazard", "VERIFIED IN CODE"),
    ("Second-Look Resolution", "Dual-spectral recheck confirming genuine/debris", "VERIFIED IN CODE"),
    ("Multi-Modal IoU Fusion", "61 assertions in aiPipeline.test.ts", "VERIFIED IN CODE"),
    ("AI Latency Benchmarking", "Measured in milliseconds via performance.now()", "VERIFIED IN CODE"),
    ("Hardware Provider Switch", "providerFactory.ts reading SIMULATION_MODE", "READY FOR PHASE 4")
]

for row_idx, (m, v, s) in enumerate(val_rows, start=1):
    c0 = tbl_val.cell(row_idx, 0)
    c1 = tbl_val.cell(row_idx, 1)
    c2 = tbl_val.cell(row_idx, 2)
    
    c0.width = Inches(2.2)
    c1.width = Inches(2.8)
    c2.width = Inches(1.7)
    
    if row_idx % 2 == 0:
        set_cell_shading(c0, "f8fafc")
        set_cell_shading(c1, "f8fafc")
        set_cell_shading(c2, "f8fafc")
        
    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_before = Pt(2)
    p0.paragraph_format.space_after = Pt(2)
    p0.add_run(m).font.size = Pt(8)
    
    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after = Pt(2)
    p1.add_run(v).font.size = Pt(8)
    
    p2 = c2.paragraphs[0]
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(s)
    r2.font.size = Pt(8)
    r2.bold = True
    r2.font.color.rgb = COLOR_GREEN

doc.save(doc_path)
print(f"Successfully generated {doc_path} ({os.path.getsize(doc_path) / 1024:.1f} KB)")
